"""
Fiche-Version Generator — Streamlit App
Run: streamlit run app.py
"""

import io
import re
import json
import time
import copy
import html as html_module
from collections import defaultdict
from datetime import datetime
from lxml import etree
import streamlit as st
from mistralai import Mistral
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Page Config ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Fiche-Version Generator",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─── CSS ──────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
  /* Global */
  .block-container { padding-top: 1.5rem; }
  h1 { color: #1F497D; }

  /* Ticket cards */
  .card {
    border: 1px solid rgba(128,128,128,0.25);
    border-radius: 10px;
    padding: 14px 18px;
    margin-bottom: 10px;
    background: var(--secondary-background-color);
    color: var(--text-color);
  }
  .card.relevant { border-left: 5px solid #2e7d32; }
  .card.skipped  { border-left: 5px solid #c62828; opacity: 0.85; }

  /* Badges */
  .badge {
    display: inline-block;
    padding: 2px 10px;
    border-radius: 12px;
    font-size: 11px;
    font-weight: 700;
    margin-right: 4px;
  }
  .badge-green  { background: rgba(46,125,50,0.15);   color: #4caf50; }
  .badge-red    { background: rgba(198,40,40,0.15);   color: #ef5350; }
  .badge-blue   { background: rgba(21,101,192,0.15);  color: #42a5f5; }
  .badge-grey   { background: rgba(128,128,128,0.15); color: var(--text-color); }
  .badge-orange { background: rgba(230,81,0,0.15);    color: #ffa726; }

  /* Sections inside card */
  .release-note {
    background: rgba(46,125,50,0.1);
    border-left: 3px solid #558b2f;
    padding: 10px 14px;
    border-radius: 5px;
    margin-top: 10px;
    font-size: 14px;
    color: var(--text-color);
  }
  .reasoning-box {
    background: rgba(128,128,128,0.08);
    border-left: 3px solid #90a4ae;
    padding: 10px 14px;
    border-radius: 5px;
    margin-top: 8px;
    font-size: 13px;
    color: var(--text-color);
    opacity: 0.8;
    font-style: italic;
  }
  .skip-box {
    background: rgba(198,40,40,0.08);
    border-left: 3px solid #ef9a9a;
    padding: 10px 14px;
    border-radius: 5px;
    margin-top: 10px;
    font-size: 13px;
    color: #ef5350;
  }
  .card-title   { font-weight: 700; font-size: 15px; margin-bottom: 4px; color: var(--text-color); }
  .summary-text { color: var(--text-color); opacity: 0.75; font-size: 13px; margin-top: 4px; }
  .label-sm     { font-size: 10px; font-weight: 600; color: var(--text-color); opacity: 0.55; text-transform: uppercase; letter-spacing: 0.5px; }

  /* Section divider */
  .section-header {
    font-size: 18px; font-weight: 700; color: #4a90d9;
    border-bottom: 2px solid rgba(74,144,217,0.3);
    padding-bottom: 6px; margin: 20px 0 14px 0;
  }

  /* Export bar */
  .export-bar {
    background: rgba(33,150,243,0.08);
    border: 1px solid rgba(33,150,243,0.25);
    border-radius: 8px;
    padding: 14px 18px;
    margin: 16px 0;
    color: var(--text-color);
  }

  /* Stats strip */
  .stat-box {
    background: var(--secondary-background-color);
    border-radius: 8px;
    padding: 10px 16px; text-align: center;
  }
  .stat-num   { font-size: 28px; font-weight: 700; color: #4a90d9; }
  .stat-label { font-size: 12px; color: var(--text-color); opacity: 0.6; }

  div[data-testid="stProgress"] > div { border-radius: 8px; }
</style>
""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════════════════════
# PROMPTS
# ═══════════════════════════════════════════════════════════════════════════════
PROMPT_LONG = SYSTEM_PROMPT_FR = """\

Tu es un rédacteur technique expert en fiches de version logicielles pour le système GID (ERP de l'administration publique).

TES DEUX SEULES TÂCHES :

1. Déterminer si le ticket décrit un changement réel et livré, pertinent pour la fiche de version.
2. Si pertinent, rédiger une phrase de note de version en français formel.

=== RÈGLES DE PERTINENCE ===


PERTINENT : changement livré et visible (fonctionnalité, amélioration, correction, changement d'interface).

=== RÈGLES DE RÉDACTION (change_description) ===

- Français formel, 1-2 phrases, du point de vue de l'utilisateur final.
- Commence toujours par un verbe d'action nominalisé :
  Ajout de... | Correction de... | Optimisation de... | Mise à jour de...
  | Amélioration de... | Prise en charge de...
- Précise le périmètre fonctionnel concerné.
- N'utilise PAS : numéros de ticket, noms de personnes, noms d'environnement
  (prod/recette/trunk), ni détails d'implémentation technique.

=== EXEMPLES ===

Exemple 1 — PERTINENT (amélioration) :
  Résumé   : MODIFICATION DU POSTE COMPTABLE VIA L'IHM
  Composant: Référentiel
  Commentaire pertinent : 'Traiter l'exception technique lors du changement de l'acteur comptable'
  → relevant: true
  → change_description: 'Modification du poste comptable assignataire via l'IHM,
    avec gestion des dates de début et de fin de relation.'

Exemple 2 — NON PERTINENT (discussion interne) :
  Résumé   : Point architecture future API
  Description: Réunion de cadrage. Rien de livré sur cette version.
  → relevant: false
  → reason_if_not_relevant: 'Discussion interne sans livraison fonctionnelle identifiable.'

Exemple 3 — PERTINENT (redressement d'un bug) :
  Résumé   : Correction calcul montant total dans le récapitulatif
  Commentaire: 'Bug corrigé dans la formule de calcul, vérifié sur plusieurs cas'
  → relevant: true
  → change_description: 'Correction du calcul du montant total dans le récapitulatif,
    prenant désormais en compte l'ensemble des lignes.'

=== FORMAT DE RÉPONSE ===
Réponds UNIQUEMENT avec un objet JSON valide, sans explication ni markdown :
{
  "relevant": true ou false,
  "change_description": "<phrase(s) de note de version en français formel>",
  "reason_if_not_relevant": "<raison brève si non pertinent, sinon null>",
  "reasoning": "<2-3 phrases : quel champ a orienté la décision, quel changement fonctionnel a été identifié, comment il a été formulé>"
}
"""

PROMPT_SMALL = """\
Tu es un rédacteur technique expert en fiches de version logicielles pour le système GID (ERP de l'administration publique).

Tu reçois un ticket JIRA :
1. Évaluer la pertinence du ticket pour la fiche de version.
2. Si pertinent, rédiger une phrase de note de version.

IMPORTANT DES CHAMPS : Résumé + derniers commentaires = signal principal. Description = contexte secondaire.
Statut = à ignorer totalement.

PERTINENT : changement livré et visible (fonctionnalité, amélioration, correction, changement d'interface).
NON PERTINENT : discussion interne | tâche annulée/reportée | dette technique sans impact utilisateur | doublon | contenu trop vague.

RÉDACTION (change_description) :
- Français formel, 1-2 phrases, point de vue utilisateur final.
- Débuter par un verbe nominalisé : Ajout de… | Correction de… | Amélioration de… | Mise à jour de… | Prise en charge de…
- Préciser le périmètre fonctionnel. Aucun numéro de ticket, nom de personne, environnement ou détail technique.

RÉPONSE — JSON valide uniquement, sans markdown :
{
  "relevant": true | false,
  "change_description": "<note de version en français>",
  "reason_if_not_relevant": "<raison brève | null>",
  "reasoning": "<2-3 phrases : quel champ a orienté la décision, quel changement a été identifié, comment il a été formulé>"
}"""


# ═══════════════════════════════════════════════════════════════════════════════
# PREPROCESSING
# ═══════════════════════════════════════════════════════════════════════════════
SELECTED_TAGS = [
    ('key',         'key'),
    ('type',        'type'),
    ('summary',     'summary'),
    ('description', 'description'),
    ('component',   'component'),
    ('comments',    'comments'),
]

SENSITIVE_KEYWORDS = [
    'recette', 'déploiement', 'déployé en', 'branch', 'trunk version',
    'salam,', 'bonjour,', 'envoyé :', 'cordialement,', 'objet :',
    'test ok', 'oui, effectivement', 'voir le document ci-joint', 'me contacter',
    'voir retour de', 'suite à notre discussion',
]

CLEAN_PATTERNS = [
    r'https?://\S+',
    r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}',
    r'\bSGID-\d+\b',
    r'\bREC-\d+\b',
    r'MC\.[a-zA-Z.]+',
    r'\[.*?\]',
    r'(?:login|username|utilisateur|user)\s*(?:comptable|ordonnateur)?\s*:?\s*[a-zA-Z]+\.[a-zA-Z]+[^\w]*',
    r'\b[A-Za-z]+\.[A-Za-z]+\b',
    r'(\d+)/(\d{4})\b',
    r'\b\d+(?:[\s.,]?\d+)*(?:\s*(?:MAD|DHs|dhs?|\$))?\b',
]

ANOMALY_TYPES = {'bug', 'bogue'}
FIELDS_TO_CLEAN = ['description', 'comments', 'summary']


def strip_html(raw: str) -> str:
    decoded = html_module.unescape(raw or '')
    clean = re.sub(r'<[^>]+>', ' ', decoded)
    return re.sub(r'\s+', ' ', clean).strip()


def extract_comments(item_el) -> str:
    block = item_el.find('comments')
    if block is None:
        return ''
    parts = [strip_html(c.text or '') for c in block.findall('comment') if c.text]
    return ' | '.join(filter(None, parts))


def parse_xml_bytes(content: bytes) -> list[dict]:
    root = etree.fromstring(content)
    tickets = []
    for item in root.iter('item'):
        ticket = {}
        for field_name, xml_tag in SELECTED_TAGS:
            if field_name == 'comments':
                ticket['comments'] = extract_comments(item)
            else:
                el = item.find(xml_tag)
                raw = (el.text or '') if el is not None else ''
                ticket[field_name] = strip_html(raw)
        tickets.append(ticket)
    return tickets


def clean_text(text: str) -> str:
    if not text:
        return text
    for pattern in CLEAN_PATTERNS:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE)
    separators = r'(?<=[.!?])\s+|\s*\|\s*|\n+'
    segments = re.split(separators, text)
    filtered = []
    for seg in segments:
        seg = seg.strip()
        if not seg:
            continue
        if any(kw in seg.lower() for kw in SENSITIVE_KEYWORDS):
            continue
        filtered.append(seg)
    result = ' | '.join(filtered) if '|' in text else ' '.join(filtered)
    return re.sub(r'\s+', ' ', result).strip()


def clean_ticket(ticket: dict) -> dict:
    cleaned = copy.deepcopy(ticket)
    for field in FIELDS_TO_CLEAN:
        cleaned[field] = clean_text(cleaned.get(field, ''))
    return cleaned


# ═══════════════════════════════════════════════════════════════════════════════
# LLM
# ═══════════════════════════════════════════════════════════════════════════════
def build_user_message(ticket: dict) -> str:
    return (
        f"Analyse ce ticket JIRA et produis une entrée de note de version.\n\n"
        f"Clé          : {ticket.get('key', '')}\n"
        f"Type         : {ticket.get('type', '')}\n"
        f"Composant    : {ticket.get('component', '')}\n\n"
        f"Résumé :\n{ticket.get('summary', '')}\n\n"
        f"Description :\n{ticket.get('description', '')[:1200]}\n\n"
        f"Commentaires pertinents :\n{ticket.get('comments', '')[:2500]}"
    )


def analyze_ticket(client: Mistral, model: str, system_prompt: str,
                   ticket: dict, delay: float, retries: int = 4) -> dict:
    messages = [
        {'role': 'system', 'content': system_prompt},
        {'role': 'user',   'content': build_user_message(ticket)},
    ]
    for attempt in range(retries):
        try:
            response = client.chat.complete(
                model=model,
                messages=messages,
                response_format={'type': 'json_object'},
                max_tokens=512,
                temperature=0.2,
            )
            raw = response.choices[0].message.content.strip()
            if raw.startswith('```'):
                raw = raw.split('```')[1]
                if raw.startswith('json'):
                    raw = raw[4:]
            result = json.loads(raw)
            result['key']       = ticket.get('key', '')
            result['component'] = ticket.get('component', '') or 'Autres évolutions'
            result['type']      = ticket.get('type', '')
            return result
        except Exception as e:
            err = str(e)
            is_rate = '429' in err
            wait = min(60, 2 ** (attempt + 2)) if is_rate else 2 ** attempt
            time.sleep(wait)
    return {
        'key': ticket.get('key', ''), 'relevant': False,
        'change_description': None, 'reasoning': None,
        'component': ticket.get('component', '') or 'Autres évolutions',
        'type': ticket.get('type', ''),
        'reason_if_not_relevant': 'Appel Mistral échoué après plusieurs tentatives.',
    }


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX GENERATION
# ═══════════════════════════════════════════════════════════════════════════════
def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_component_section(doc, component_name: str, changes: list):
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.add_run(f'Module « {component_name} »')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    intro = doc.add_paragraph('Ce module intègre les évolutions suivantes dans cette version :')
    intro.paragraph_format.space_before = Pt(0)
    intro.paragraph_format.space_after  = Pt(2)

    for change in changes:
        bullet = doc.add_paragraph(style='List Bullet')
        bullet.add_run(change)
        bullet.paragraph_format.space_after = Pt(2)

    doc.add_paragraph('').paragraph_format.space_after = Pt(4)


def generate_docx_bytes(selected_results: list, version_label: str, deploy_date: str) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Title table
    tt = doc.add_table(rows=2, cols=2)
    tt.style = 'Table Grid'
    c_date  = tt.cell(0, 0);  c_date.text  = f'Date : {deploy_date}'
    c_title = tt.cell(0, 1);  c_title.text = 'Fiche de version GID'
    set_cell_bg(c_date,  '1F497D')
    set_cell_bg(c_title, '1F497D')
    for cell in [c_date, c_title]:
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True
            run.font.size = Pt(13)
    desc_cell = tt.cell(1, 0)
    desc_cell.merge(tt.cell(1, 1))
    desc_cell.text = (
        f'La version {version_label} déployée en production le {deploy_date} '
        f'intègre les fonctionnalités suivantes :'
    )
    set_cell_bg(desc_cell, 'D0E4F7')
    doc.add_paragraph('')

    # Group by component / anomaly
    component_changes: dict = defaultdict(list)
    anomaly_changes: list   = []
    for r in selected_results:
        is_anom = r.get('type', '').lower() in ANOMALY_TYPES
        if is_anom:
            anomaly_changes.append((r['component'], r['change_description']))
        else:
            component_changes[r['component']].append(r['change_description'])

    for comp in sorted(component_changes.keys()):
        add_component_section(doc, comp, component_changes[comp])

    if anomaly_changes:
        ah = doc.add_paragraph()
        run = ah.add_run('Anomalies :')
        run.bold = True; run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        intro = doc.add_paragraph('Cette version corrige également les anomalies se rapportant à :')
        intro.paragraph_format.space_after = Pt(2)
        for comp, change in anomaly_changes:
            b = doc.add_paragraph(style='List Bullet')
            b.add_run(f'[{comp}] {change}' if comp else change)
            b.paragraph_format.space_after = Pt(2)

    # Footer
    doc.add_paragraph('')
    fp = doc.add_paragraph(
        f'Document généré automatiquement le {datetime.now().strftime("%d/%m/%Y à %H:%M")}'
    )
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.runs[0]
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    fr.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# CARD RENDERER
# ═══════════════════════════════════════════════════════════════════════════════
def render_card(result: dict):
    key       = result.get('key', '?')
    relevant  = result.get('relevant', False)
    component = result.get('component', '—')
    ticket_type = result.get('type', '—')
    change    = result.get('change_description') or ''
    reasoning = result.get('reasoning') or ''
    skip_rsn  = result.get('reason_if_not_relevant') or ''
    is_anom   = ticket_type.lower() in ANOMALY_TYPES

    card_class  = 'card relevant' if relevant else 'card skipped'
    badge_rel   = '<span class="badge badge-green">✅ PERTINENT</span>' if relevant else '<span class="badge badge-red">⏭ IGNORÉ</span>'
    badge_comp  = f'<span class="badge badge-blue">{component}</span>'
    badge_type  = f'<span class="badge badge-orange">{ticket_type}</span>' if is_anom else f'<span class="badge badge-grey">{ticket_type}</span>'

    content_html = ''
    if relevant and change:
        content_html += f'<div class="release-note"><span class="label-sm">Note de version</span><br>{change}</div>'
    if not relevant and skip_rsn:
        content_html += f'<div class="skip-box"><span class="label-sm">Raison</span><br>{skip_rsn}</div>'
    if reasoning:
        content_html += f'<div class="reasoning-box"><span class="label-sm">Raisonnement du modèle</span><br>{reasoning}</div>'

    st.markdown(f"""
    <div class="{card_class}">
      <div class="card-title">{key} &nbsp; {badge_rel} &nbsp; {badge_comp} &nbsp; {badge_type}</div>
      {content_html}
    </div>
    """, unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════════════════════
# SESSION STATE INIT
# ═══════════════════════════════════════════════════════════════════════════════
for key in ('all_results', 'processed', 'cleaned_tickets'):
    if key not in st.session_state:
        st.session_state[key] = [] if key != 'processed' else False


# ═══════════════════════════════════════════════════════════════════════════════
# SIDEBAR
# ═══════════════════════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("## ⚙️ Configuration")
    st.divider()

    api_key = st.text_input(
        "Clé API Mistral",
        type="password",
        placeholder="Votre clé depuis console.mistral.ai",
        help="Clé gratuite sur console.mistral.ai"
    )

    model_choice = st.selectbox(
        "Modèle",
        options=['mistral-small-latest', 'mistral-medium-latest', 'open-mistral-7b', 'open-mixtral-8x7b'],
        index=0,
    )

    prompt_choice = st.radio(
        "Version du prompt",
        options=["🔹 Court (moins de tokens)", "🔷 Long (plus de contexte)"],
        index=0,
        help="Le prompt long inclut des exemples few-shot et des règles détaillées. Le prompt court est plus rapide."
    )
    selected_prompt = PROMPT_LONG if "Long" in prompt_choice else PROMPT_SMALL

    st.divider()
    st.markdown("### 📄 Document")

    version_label = st.text_input("Label de version", value="Trunk")
    deploy_date   = st.text_input("Date de déploiement", value=datetime.now().strftime("%d/%m/%Y"))
    inter_delay   = 1.2  # Delay between API calls to mitigate rate limits



# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════
st.markdown("# 📋 Fiche-Version Generator")
st.markdown("Importez vos tickets JIRA en XML, analysez-les avec Mistral, sélectionnez et exportez.")
st.divider()

# ── Step 1: Upload ─────────────────────────────────────────────────────────────
st.markdown('<div class="section-header">① Importer les tickets XML</div>', unsafe_allow_html=True)

uploaded_files = st.file_uploader(
    "Glissez-déposez un ou plusieurs fichiers XML JIRA",
    type=["xml"],
    accept_multiple_files=True,
    help="Export les Tickets JIRA au format XML — un ou plusieurs fichiers"
)

if uploaded_files:
    st.success(f"✅ **{len(uploaded_files)}** fichier(s) prêts à être traités.")

# ── Step 2: Process ────────────────────────────────────────────────────────────
st.markdown('<div class="section-header">② Analyser avec Mistral</div>', unsafe_allow_html=True)

col_btn, col_info = st.columns([2, 5])
with col_btn:
    run_btn = st.button(
        "▶ Analyser les tickets",
        type="primary",
        disabled=not (uploaded_files and api_key),
        use_container_width=True,
    )
with col_info:
    if not api_key:
        st.warning("⚠️ Renseignez votre clé API Mistral dans la barre latérale.")
    elif not uploaded_files:
        st.info("📂 Importez au moins un fichier XML.")
    else:
        n = len(uploaded_files)
        est = int(n * inter_delay) + n * 2
        st.info(f"Prêt · {n} fichier(s) · durée estimée ~{est}s à {inter_delay}s/appel")

if run_btn and uploaded_files and api_key:
    # Parse + clean
    raw_tickets = []
    for f in uploaded_files:
        try:
            parsed = parse_xml_bytes(f.read())
            raw_tickets.extend(parsed)
        except Exception as e:
            st.error(f"❌ Erreur parsing `{f.name}`: {e}")

    cleaned = [clean_ticket(t) for t in raw_tickets]

    # ── Deduplicate by ticket key (keep first occurrence) ──────────────────
    seen_keys: set = set()
    deduped, duplicates = [], []
    for t in cleaned:
        k = t.get('key', '')
        if k and k in seen_keys:
            duplicates.append(k)
        else:
            if k:
                seen_keys.add(k)
            deduped.append(t)
    cleaned = deduped
    if duplicates:
        st.warning(
            f"⚠️ **{len(duplicates)}** ticket(s) en doublon ignoré(s) "
            f"(première occurrence conservée) : {', '.join(duplicates)}"
        )

    st.session_state.cleaned_tickets = cleaned

    if not cleaned:
        st.error("Aucun ticket valide trouvé dans les fichiers.")
    else:
        client = Mistral(api_key=api_key)
        all_results = []
        progress_bar = st.progress(0, text="Initialisation…")
        status_area  = st.empty()

        for i, ticket in enumerate(cleaned):
            status_area.markdown(
                f"🔄 Analyse de **{ticket.get('key', '?')}** "
                f"[{ticket.get('component', '')}] · {i+1}/{len(cleaned)}"
            )
            result = analyze_ticket(client, model_choice, selected_prompt, ticket, inter_delay)
            all_results.append(result)

            # Initialize checkbox state as True for relevant, False for skipped
            cb_key = f"sel_{result['key']}"
            if cb_key not in st.session_state:
                st.session_state[cb_key] = result.get('relevant', False)

            pct = (i + 1) / len(cleaned)
            progress_bar.progress(pct, text=f"Traités : {i+1}/{len(cleaned)}")
            if i < len(cleaned) - 1:
                time.sleep(inter_delay)

        st.session_state.all_results = all_results
        st.session_state.processed   = True
        progress_bar.empty()
        status_area.empty()
        st.success(f"✅ Analyse terminée — {sum(1 for r in all_results if r.get('relevant'))} pertinents · "
                   f"{sum(1 for r in all_results if not r.get('relevant'))} ignorés")
        st.rerun()


# ── Step 3: Results ────────────────────────────────────────────────────────────
if st.session_state.get('processed') and st.session_state.get('all_results'):
    all_results = st.session_state.all_results
    relevant = [r for r in all_results if r.get('relevant')]
    skipped  = [r for r in all_results if not r.get('relevant')]

    # Stats strip
    st.markdown('<div class="section-header">③ Résultats et sélection</div>', unsafe_allow_html=True)
    s1, s2, s3, s4 = st.columns(4)
    with s1:
        st.markdown(f'<div class="stat-box"><div class="stat-num">{len(all_results)}</div><div class="stat-label">Tickets analysés</div></div>', unsafe_allow_html=True)
    with s2:
        st.markdown(f'<div class="stat-box"><div class="stat-num" style="color:#2e7d32">{len(relevant)}</div><div class="stat-label">Pertinents</div></div>', unsafe_allow_html=True)
    with s3:
        st.markdown(f'<div class="stat-box"><div class="stat-num" style="color:#c62828">{len(skipped)}</div><div class="stat-label">Ignorés</div></div>', unsafe_allow_html=True)
    components = sorted(set(r['component'] for r in relevant))
    with s4:
        st.markdown(f'<div class="stat-box"><div class="stat-num">{len(components)}</div><div class="stat-label">Composants</div></div>', unsafe_allow_html=True)

    st.markdown("")

    # Select all / deselect controls
    ca_col, _, filter_col = st.columns([2, 5, 3])
    with ca_col:
        select_all = st.checkbox("Tout sélectionner", value=False)
        if select_all != st.session_state.get("_select_all_prev", False):
            st.session_state["_select_all_prev"] = select_all
            for r in all_results:
                st.session_state[f"sel_{r['key']}"] = select_all
            st.rerun()
    with filter_col:
        show_skipped = st.toggle("Afficher les tickets ignorés", value=False)

    st.markdown("")

    # Render tickets grouped by component
    display_results = all_results if show_skipped else relevant

    # Group by component for display
    by_component: dict = defaultdict(list)
    for r in display_results:
        by_component[r['component']].append(r)

    # Relevant components first, then "skipped only" components
    for comp in sorted(by_component.keys()):
        comp_results = by_component[comp]
        with st.expander(f"**{comp}** — {len(comp_results)} ticket(s)", expanded=True):
            for result in comp_results:
                key = result.get('key', '?')
                cb_col, card_col = st.columns([1, 20])
                with cb_col:
                    st.checkbox(
                        "",
                        key=f"sel_{key}",
                        #help="Cocher pour inclure dans l'export",
                    )
                with card_col:
                    render_card(result)

    # ── Step 4: Export ─────────────────────────────────────────────────────────
    st.markdown('<div class="section-header">④ Exporter la Fiche de Version</div>', unsafe_allow_html=True)

    selected_results = [
        r for r in all_results
        if st.session_state.get(f"sel_{r.get('key', '')}", False)
        and r.get('relevant') and r.get('change_description')
    ]

    st.markdown(f"""
    <div class="export-bar">
      <strong>{len(selected_results)}</strong> ticket(s) sélectionné(s) pour l'export
      · groupés en <strong>{len(set(r['component'] for r in selected_results))}</strong> composant(s)
    </div>
    """, unsafe_allow_html=True)

    if selected_results:
        docx_bytes = generate_docx_bytes(selected_results, version_label, deploy_date)
        filename   = f"Fiche_Version_GID_{version_label.replace(' ', '_')}_{deploy_date.replace('/', '-')}.docx"

        st.download_button(
            label="⬇ Télécharger la Fiche de Version (.docx)",
            data=docx_bytes,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            use_container_width=False,
        )
        st.caption(f"Fichier : `{filename}` · {len(selected_results)} entrées · {len(set(r['component'] for r in selected_results))} modules")
    else:
        st.info("Sélectionnez au moins un ticket pertinent pour générer le document.")